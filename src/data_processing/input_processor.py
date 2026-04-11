"""
Complete Input Processor
Handles all input types with production-ready processing
"""

import os
import re
import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
from datetime import datetime

# Setup logging — no heavy libs at module level (keeps import instant)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
# All heavy libs (whisper, easyocr, pdfplumber, cv2, pandas, ftfy, docx,
# requests, newspaper, LanguageDetector) are imported lazily inside the
# methods that need them so InputProcessor() instantiates in <1s.


class InputProcessor:
    """
    Production-ready input processor for all input types.
    All heavy models are lazy-loaded on first use — __init__ is instant.
    """

    def __init__(
        self,
        whisper_model_size: str = "base",
        ocr_languages: List[str] = None
    ):
        """
        Initialize input processor (instantaneous — no models loaded here).

        Args:
            whisper_model_size: Whisper model size (tiny, base, small, medium, large)
            ocr_languages: Languages for OCR (default: ['en', 'hi'])
        """
        # Store config — actual models loaded on first use
        self._whisper_model_size = whisper_model_size
        self._ocr_languages      = ocr_languages or ['en', 'hi']

        # Lazy handles (populated on first access)
        self._whisper_model   = None
        self._ocr_reader      = None
        self._language_detector = None

        logger.info("✓ InputProcessor initialized successfully")

    # ── Lazy accessors ────────────────────────────────────────────────────────

    @property
    def whisper_model(self):
        if self._whisper_model is None:
            import whisper as _whisper
            logger.info(f"Loading Whisper model ({self._whisper_model_size})...")
            self._whisper_model = _whisper.load_model(self._whisper_model_size)
            logger.info("✓ Whisper model loaded")
        return self._whisper_model

    @property
    def ocr_reader(self):
        if self._ocr_reader is None:
            import easyocr
            logger.info(f"Loading EasyOCR for languages: {self._ocr_languages}...")
            try:
                self._ocr_reader = easyocr.Reader(self._ocr_languages, gpu=False)
            except ValueError as e:
                logger.warning(f"OCR language combination failed: {e}. Falling back to English-only.")
                self._ocr_reader = easyocr.Reader(['en'], gpu=False)
            logger.info("✓ EasyOCR loaded")
        return self._ocr_reader

    @property
    def language_detector(self):
        if self._language_detector is None:
            from src.preprocessing.language_detector import LanguageDetector
            logger.info("Loading FastText language detector...")
            self._language_detector = LanguageDetector()
            logger.info("✓ Language detector loaded")
        return self._language_detector

    def process(
        self,
        input_data: Union[str, bytes, Path],
        input_type: str,
        metadata: Optional[Dict] = None
    ) -> Dict[str, Any]:
        """
        Main processing function

        Args:
            input_data: The input (text, file path, URL, bytes)
            input_type: One of ['text', 'voice', 'pdf', 'docx', 'image', 'url']
            metadata: Optional metadata to include

        Returns:
            {
                'text': str,              # Extracted clean text
                'language': str,          # Detected language code
                'source_type': str,       # Input type
                'metadata': dict,         # Additional info
                'processing_info': dict   # Processing details
            }
        """
        logger.info(f"Processing input of type: {input_type}")

        processors = {
            'text': self._process_text,
            'voice': self._process_voice,
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'xlsx': self._process_xlsx,
            'csv': self._process_csv,
            'image': self._process_image,
            'url': self._process_url
        }

        if input_type not in processors:
            raise ValueError(f"Unsupported input type: {input_type}. "
                           f"Supported: {list(processors.keys())}")

        # Validate input
        self._validate_input(input_data, input_type)

        # Process
        start_time = datetime.now()
        result = processors[input_type](input_data)
        processing_time = (datetime.now() - start_time).total_seconds()

        # Add processing info
        result['processing_info'] = {
            'processing_time_seconds': processing_time,
            'processed_at': datetime.now().isoformat(),
            'processor_version': '1.0'
        }

        # Add user metadata if provided
        if metadata:
            result['metadata'].update(metadata)

        logger.info(f"✓ Processing complete in {processing_time:.2f}s")

        return result

    # ==========================================
    # TEXT INPUT
    # ==========================================

    def _process_text(self, text: str) -> Dict[str, Any]:
        """
        Process direct text input

        Steps:
        1. Fix encoding issues
        2. Clean text
        3. Detect language
        4. Validate length
        """
        logger.info("Processing text input...")

        # Step 1: Fix encoding
        try:
            import ftfy
            text = ftfy.fix_text(text)
        except ImportError:
            pass  # ftfy not installed, skip encoding fix

        # Step 2: Clean text
        text_cleaned = self._clean_text(text)

        # Step 3: Detect language
        language = self._detect_language(text_cleaned)

        # Step 4: Validate
        if len(text_cleaned) < 10:
            logger.warning("Text is very short (< 10 characters)")

        return {
            'text': text_cleaned,
            'language': language,
            'source_type': 'text',
            'metadata': {
                'original_length': len(text),
                'cleaned_length': len(text_cleaned),
                'num_words': len(text_cleaned.split()),
                'num_sentences': len(self._split_sentences(text_cleaned))
            }
        }

    # ==========================================
    # VOICE INPUT
    # ==========================================

    def _process_voice(self, audio_path: str) -> Dict[str, Any]:
        """
        Process voice/audio input using Whisper

        Steps:
        1. Validate audio file
        2. Transcribe using Whisper
        3. Extract language from Whisper
        4. Clean transcript

        Supported formats: .wav, .mp3, .m4a, .flac, .ogg
        """
        logger.info(f"Processing voice input from: {audio_path}")

        # Step 1: Validate audio file
        if not os.path.exists(audio_path):
            raise FileNotFoundError(f"Audio file not found: {audio_path}")

        # Get audio metadata
        try:
            import soundfile as sf
            audio_data, sample_rate = sf.read(audio_path)
            duration = len(audio_data) / sample_rate
        except ImportError:
            raise ImportError("pip install soundfile")
        except Exception as e:
            raise ValueError(f"Failed to read audio file: {str(e)}")

        logger.info(f"Audio duration: {duration:.2f} seconds")

        # Step 2: Transcribe using Whisper
        logger.info("Transcribing audio with Whisper...")
        result = self.whisper_model.transcribe(
            audio_path,
            language=None,  # Auto-detect
            task="transcribe",
            verbose=False
        )

        # Step 3: Extract information
        text = result['text'].strip()
        language = result['language']

        # Whisper also provides word-level timestamps
        segments = result.get('segments', [])

        # Step 4: Clean transcript
        text_cleaned = self._clean_text(text)

        logger.info(f"Transcription complete. Language: {language}")

        return {
            'text': text_cleaned,
            'language': language,
            'source_type': 'voice',
            'metadata': {
                'audio_path': audio_path,
                'duration_seconds': duration,
                'sample_rate': sample_rate,
                'num_segments': len(segments),
                'original_transcript': text,
                'whisper_confidence': self._calculate_whisper_confidence(segments)
            }
        }

    def _calculate_whisper_confidence(self, segments: List[Dict]) -> float:
        """Calculate average confidence from Whisper segments"""
        if not segments:
            return 1.0

        # Whisper doesn't provide confidence directly,
        # but we can estimate from segment properties
        # For now, return 1.0 (placeholder)
        return 1.0

    # ==========================================
    # PDF INPUT
    # ==========================================

    def _process_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """
        Process PDF document

        Steps:
        1. Extract text from PDF
        2. Extract tables
        3. Handle scanned PDFs (OCR)
        4. Merge all content
        5. Detect language

        Uses pdfplumber for better extraction than PyPDF2
        """
        logger.info(f"Processing PDF: {pdf_path}")

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")

        full_text = ""
        tables = []
        num_pages = 0
        is_scanned = False

        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                num_pages = len(pdf.pages)
                logger.info(f"PDF has {num_pages} pages")

                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()

                    if page_text:
                        full_text += page_text + "\n"
                    else:
                        # Page might be scanned - use OCR
                        logger.info(f"Page {page_num} appears scanned, using OCR...")
                        is_scanned = True

                        # Convert page to image
                        page_image = page.to_image(resolution=300)
                        img_bytes = page_image.original

                        # OCR
                        ocr_text = self._ocr_image(img_bytes)
                        full_text += ocr_text + "\n"

                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
                        logger.info(f"Found {len(page_tables)} tables on page {page_num}")

        except Exception as e:
            raise ValueError(f"Failed to process PDF: {str(e)}")

        # Clean text
        text_cleaned = self._clean_text(full_text)

        # Detect language
        language = self._detect_language(text_cleaned)

        # Format tables as text
        tables_text = self._format_tables(tables)

        # Combine text and tables
        combined_text = text_cleaned
        if tables_text:
            combined_text += "\n\nTables:\n" + tables_text

        return {
            'text': combined_text,
            'language': language,
            'source_type': 'pdf',
            'metadata': {
                'file_path': pdf_path,
                'num_pages': num_pages,
                'is_scanned': is_scanned,
                'num_tables': len(tables),
                'text_length': len(text_cleaned),
                'tables_data': tables[:5]  # Store first 5 tables
            }
        }

    def _format_tables(self, tables: List[List[List]]) -> str:
        """Format extracted tables as readable text"""
        if not tables:
            return ""

        formatted = []
        for i, table in enumerate(tables, 1):
            formatted.append(f"Table {i}:")
            for row in table:
                # Filter out None values
                row_cleaned = [str(cell) if cell else "" for cell in row]
                formatted.append(" | ".join(row_cleaned))
            formatted.append("")  # Empty line between tables

        return "\n".join(formatted)

    # ==========================================
    # DOCX INPUT
    # ==========================================

    def _process_docx(self, docx_path: str) -> Dict[str, Any]:
        """
        Process Word document (.docx)

        Steps:
        1. Extract paragraphs
        2. Extract tables
        3. Preserve document structure
        4. Detect language
        """
        logger.info(f"Processing DOCX: {docx_path}")

        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        try:
            from docx import Document
            doc = Document(docx_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Combine text
            full_text = "\n\n".join(paragraphs)

            # Add tables
            if tables:
                tables_text = self._format_tables(tables)
                full_text += "\n\n" + tables_text

            # Clean
            text_cleaned = self._clean_text(full_text)

            # Detect language
            language = self._detect_language(text_cleaned)

        except Exception as e:
            raise ValueError(f"Failed to process DOCX: {str(e)}")

        return {
            'text': text_cleaned,
            'language': language,
            'source_type': 'docx',
            'metadata': {
                'file_path': docx_path,
                'num_paragraphs': len(paragraphs),
                'num_tables': len(tables),
                'text_length': len(text_cleaned)
            }
        }

    # ==========================================
    # EXCEL/XLSX INPUT
    # ==========================================

    def _process_xlsx(self, xlsx_path: str) -> Dict[str, Any]:
        """
        Process Excel spreadsheet (.xlsx, .xls)

        Steps:
        1. Read all sheets
        2. Convert to text representation
        3. Extract key statistics
        4. Detect language from text cells
        """
        logger.info(f"Processing XLSX: {xlsx_path}")

        if not os.path.exists(xlsx_path):
            raise FileNotFoundError(f"Excel file not found: {xlsx_path}")

        try:
            # Read all sheets
            import pandas as pd
            excel_file = pd.ExcelFile(xlsx_path)
            sheet_names = excel_file.sheet_names
            
            logger.info(f"Found {len(sheet_names)} sheets: {sheet_names}")

            all_text = []
            all_sheets_data = []

            for sheet_name in sheet_names:
                # Read sheet
                df = pd.read_excel(xlsx_path, sheet_name=sheet_name)
                
                # Convert to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False)
                all_text.append(sheet_text)
                
                # Store sheet info
                all_sheets_data.append({
                    'sheet_name': sheet_name,
                    'num_rows': len(df),
                    'num_columns': len(df.columns),
                    'columns': df.columns.tolist()
                })

            # Combine all sheets
            full_text = "\n\n" + "="*50 + "\n\n".join(all_text)

            # Clean
            text_cleaned = self._clean_text(full_text)

            # Detect language
            language = self._detect_language(text_cleaned)

        except Exception as e:
            raise ValueError(f"Failed to process Excel file: {str(e)}")

        return {
            'text': text_cleaned,
            'language': language,
            'source_type': 'xlsx',
            'metadata': {
                'file_path': xlsx_path,
                'num_sheets': len(sheet_names),
                'sheet_names': sheet_names,
                'sheets_data': all_sheets_data,
                'text_length': len(text_cleaned)
            }
        }

    # ==========================================
    # CSV INPUT
    # ==========================================

    def _process_csv(self, csv_path: str) -> Dict[str, Any]:
        """
        Process CSV file

        Steps:
        1. Read CSV
        2. Convert to text representation
        3. Extract statistics
        4. Detect language
        """
        logger.info(f"Processing CSV: {csv_path}")

        if not os.path.exists(csv_path):
            raise FileNotFoundError(f"CSV file not found: {csv_path}")

        try:
            # Read CSV
            import pandas as pd
            df = pd.read_csv(csv_path)

            # Convert to text
            full_text = df.to_string(index=False)

            # Clean
            text_cleaned = self._clean_text(full_text)

            # Detect language
            language = self._detect_language(text_cleaned)

        except Exception as e:
            raise ValueError(f"Failed to process CSV: {str(e)}")

        return {
            'text': text_cleaned,
            'language': language,
            'source_type': 'csv',
            'metadata': {
                'file_path': csv_path,
                'num_rows': len(df),
                'num_columns': len(df.columns),
                'columns': df.columns.tolist(),
                'text_length': len(text_cleaned)
            }
        }

    # ==========================================
    # IMAGE INPUT (OCR)
    # ==========================================

    def _process_image(self, image_path: str) -> Dict[str, Any]:
        """
        Process image with text using OCR

        Steps:
        1. Load image
        2. Preprocess image (enhance quality)
        3. Run OCR (EasyOCR)
        4. Extract text
        5. Detect language

        Supports: .jpg, .jpeg, .png, .bmp, .tiff
        """
        logger.info(f"Processing image: {image_path}")

        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image file not found: {image_path}")

        try:
            # Load image
            import cv2
            import numpy as np
            img = cv2.imread(image_path)

            if img is None:
                raise ValueError("Failed to load image")

            # Get image dimensions
            height, width = img.shape[:2]

            # Preprocess image for better OCR
            img_processed = self._preprocess_image_for_ocr(img)

            # Run OCR
            logger.info("Running OCR...")
            ocr_results = self.ocr_reader.readtext(img_processed)

            # Extract text and confidence
            texts = []
            confidences = []

            for (bbox, text, confidence) in ocr_results:
                texts.append(text)
                confidences.append(confidence)

            # Combine all text
            full_text = " ".join(texts)

            # Clean
            text_cleaned = self._clean_text(full_text)

            # Detect language
            language = self._detect_language(text_cleaned)

            # Average confidence
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

            logger.info(f"OCR complete. Confidence: {avg_confidence:.2f}")

        except Exception as e:
            raise ValueError(f"Failed to process image: {str(e)}")

        return {
            'text': text_cleaned,
            'language': language,
            'source_type': 'image',
            'metadata': {
                'file_path': image_path,
                'image_width': width,
                'image_height': height,
                'num_text_regions': len(ocr_results),
                'avg_confidence': avg_confidence,
                'all_confidences': confidences
            }
        }

    def _preprocess_image_for_ocr(self, img: np.ndarray) -> np.ndarray:
        """
        Preprocess image to improve OCR accuracy

        Steps:
        - Convert to grayscale
        - Denoise
        - Increase contrast
        - Binarization (optional)
        """
        # Convert to grayscale
        import cv2
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray)

        # Increase contrast (CLAHE)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        contrast = clahe.apply(denoised)

        # Optional: Binarization (for very clear text)
        # _, binary = cv2.threshold(contrast, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        # return binary

        return contrast

    def _ocr_image(self, img_data: Union[np.ndarray, bytes]) -> str:
        """Run OCR on image data"""
        results = self.ocr_reader.readtext(img_data)
        text = " ".join([text for (bbox, text, conf) in results])
        return text

    # ==========================================
    # URL INPUT (Web Scraping)
    # ==========================================

    def _process_url(self, url: str) -> Dict[str, Any]:
        """
        Process article from URL

        Steps:
        1. Fetch webpage
        2. Extract article using newspaper3k
        3. Fallback to BeautifulSoup if needed
        4. Clean extracted text
        5. Detect language
        """
        logger.info(f"Processing URL: {url}")

        # Validate URL
        if not url.startswith(('http://', 'https://')):
            raise ValueError(f"Invalid URL: {url}")

        try:
            # Method 1: newspaper3k (best for news articles)
            from newspaper import Article
            article = Article(url)
            article.download()
            article.parse()

            text = article.text
            title = article.title
            authors = article.authors
            publish_date = article.publish_date

            if not text or len(text) < 50:
                # Fallback to BeautifulSoup
                logger.info("newspaper3k failed, using BeautifulSoup fallback...")
                text, title = self._scrape_with_beautifulsoup(url)
                authors = []
                publish_date = None

            # Clean text
            text_cleaned = self._clean_text(text)

            # Detect language
            language = self._detect_language(text_cleaned)

        except Exception as e:
            raise ValueError(f"Failed to scrape URL: {str(e)}")

        return {
            'text': text_cleaned,
            'language': language,
            'source_type': 'url',
            'metadata': {
                'url': url,
                'title': title,
                'authors': authors,
                'publish_date': str(publish_date) if publish_date else None,
                'text_length': len(text_cleaned),
                'num_words': len(text_cleaned.split())
            }
        }

    def _scrape_with_beautifulsoup(self, url: str) -> tuple:
        """Fallback scraping method using BeautifulSoup"""
        import requests
        from bs4 import BeautifulSoup
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.content, 'html.parser')

        # Try to find title
        title = soup.find('h1')
        title = title.get_text() if title else soup.title.string if soup.title else ""

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text, title

    # ==========================================
    # BATCH AND CONTEXT PROCESSING
    # ==========================================

    def process_with_context(
        self,
        claim_input: Union[str, Dict],
        context_documents: List[Dict],
        claim_type: str = 'text'
    ) -> Dict[str, Any]:
        """
        Process user input with both claim and supporting context documents
        
        USE CASE: User provides claim + documents to verify against
        
        Args:
            claim_input: The main claim to verify
                - If str: direct text claim
                - If dict: {'data': ..., 'type': 'text'/'voice'/etc}
                
            context_documents: List of context documents
                Each document: {'data': file_path/url/text, 'type': 'pdf'/'docx'/etc}
                
            claim_type: Type of claim input (if claim_input is str)
        
        Returns:
            {
                'claim': {
                    'text': str,
                    'language': str,
                    'source_type': str,
                    'metadata': dict
                },
                'context_documents': [
                    {
                        'text': str,
                        'language': str,
                        'source_type': str,
                        'metadata': dict,
                        'priority': 'high'  # User-provided = high priority
                    },
                    ...
                ],
                'processing_mode': 'claim_with_context'
            }
        """
        logger.info("Processing claim with user-provided context documents")
        
        # ==========================================
        # STEP 1: Process the claim
        # ==========================================
        
        if isinstance(claim_input, str):
            # Direct text claim
            claim_result = self.process(claim_input, claim_type)
        elif isinstance(claim_input, dict):
            # Structured claim input
            claim_result = self.process(
                claim_input['data'],
                claim_input['type']
            )
        else:
            raise TypeError("claim_input must be str or dict")
        
        logger.info(f"✓ Claim processed: '{claim_result['text'][:100]}...'")
        
        # ==========================================
        # STEP 2: Process context documents
        # ==========================================
        
        processed_contexts = []
        
        for i, doc in enumerate(context_documents, 1):
            logger.info(f"Processing context document {i}/{len(context_documents)}: {doc['type']}")
            
            try:
                # Process document
                doc_result = self.process(
                    doc['data'],
                    doc['type'],
                    metadata={'context_doc_index': i}
                )
                
                # Mark as high-priority context
                doc_result['priority'] = 'high'
                doc_result['is_user_provided'] = True
                
                # Add document name if provided
                if 'name' in doc:
                    doc_result['metadata']['document_name'] = doc['name']
                
                processed_contexts.append(doc_result)
                
                logger.info(f"  ✓ Context doc {i} processed: {len(doc_result['text'])} chars")
                
            except Exception as e:
                logger.error(f"  ✗ Failed to process context doc {i}: {str(e)}")
                # Continue with other documents
                continue
        
        logger.info(f"✓ Processed {len(processed_contexts)}/{len(context_documents)} context documents")
        
        # ==========================================
        # STEP 3: Combine results
        # ==========================================
        
        result = {
            'claim': claim_result,
            'context_documents': processed_contexts,
            'processing_mode': 'claim_with_context',
            'metadata': {
                'num_context_docs': len(processed_contexts),
                'context_languages': list(set(doc['language'] for doc in processed_contexts)),
                'total_context_length': sum(len(doc['text']) for doc in processed_contexts),
                'processed_at': datetime.now().isoformat()
            }
        }
        
        return result

    def process_batch(
        self,
        inputs: List[Dict]
    ) -> List[Dict[str, Any]]:
        """
        Process multiple inputs in batch
        
        Args:
            inputs: List of input dicts, each with 'data' and 'type' keys
            
            Example:
            inputs = [
                {'data': 'text claim here', 'type': 'text'},
                {'data': 'report.pdf', 'type': 'pdf'},
                {'data': 'audio.wav', 'type': 'voice'}
            ]
        
        Returns:
            List of processed results
        """
        logger.info(f"Processing batch of {len(inputs)} inputs")
        
        results = []
        
        for i, inp in enumerate(inputs, 1):
            logger.info(f"Processing input {i}/{len(inputs)}: {inp['type']}")
            
            try:
                result = self.process(
                    inp['data'],
                    inp['type'],
                    metadata=inp.get('metadata', {})
                )
                results.append(result)
                
            except Exception as e:
                logger.error(f"Failed to process input {i}: {str(e)}")
                results.append({
                    'error': str(e),
                    'input_type': inp['type'],
                    'failed': True
                })
        
        logger.info(f"✓ Batch processing complete: {len(results)} results")
        
        return results

    # ==========================================
    # UTILITY FUNCTIONS
    # ==========================================

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text

        Steps:
        - Fix encoding
        - Remove extra whitespace
        - Remove special characters (keep punctuation)
        - Normalize line breaks
        """
        # Fix encoding
        try:
            import ftfy
            text = ftfy.fix_text(text)
        except ImportError:
            pass  # ftfy not installed, skip encoding fix

        # Remove URLs
        text = re.sub(r'http\S+|www\.\S+', '', text)

        # Remove email addresses
        text = re.sub(r'\S+@\S+', '', text)

        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove leading/trailing whitespace
        text = text.strip()

        return text

    def _detect_language(self, text: str) -> str:
        """
        Detect language using langdetect (fast, <100ms, no model download).

        Falls back to Unicode script-range heuristic for very short text.

        Returns:
            ISO 639-1 language code (en, hi, ta, es, etc.)
        """
        if not text or len(text) < 5:
            return 'unknown'

        try:
            from langdetect import detect
            lang_code = detect(text)
            logger.debug(f"Detected language: {lang_code}")
            return lang_code
        except Exception:
            pass

        # Script-based fallback for very short / script-heavy text
        for char in text:
            cp = ord(char)
            if 0x0900 <= cp <= 0x097F: return 'hi'   # Devanagari
            if 0x0B80 <= cp <= 0x0BFF: return 'ta'   # Tamil
            if 0x0C00 <= cp <= 0x0C7F: return 'te'   # Telugu
            if 0x0600 <= cp <= 0x06FF: return 'ar'   # Arabic
            if 0x4E00 <= cp <= 0x9FFF: return 'zh'   # Chinese
            if 0x0400 <= cp <= 0x04FF: return 'ru'   # Cyrillic
        return 'unknown'
    
    def _detect_language_with_details(self, text: str) -> dict:
        """
        Detect language with full details
        
        Returns:
            {
                'language': str,
                'confidence': float,
                'language_name': str,
                'alternatives': [(lang, conf), ...]
            }
        """
        if not text or len(text) < 5:
            return {
                'language': 'unknown',
                'confidence': 0.0,
                'language_name': 'Unknown',
                'alternatives': []
            }
        
        try:
            lang_code, confidence = self.language_detector.detect(text)
            alternatives = self.language_detector.detect_multiple(text, top_k=3)
            
            return {
                'language': lang_code,
                'confidence': confidence,
                'language_name': self.language_detector.get_language_name(lang_code),
                'alternatives': alternatives
            }
        except Exception as e:
            logger.error(f"Language detection failed: {str(e)}")
            return {
                'language': 'unknown',
                'confidence': 0.0,
                'language_name': 'Unknown',
                'alternatives': []
            }

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences (simple version)"""
        # Simple sentence splitting
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]

    def _validate_input(self, input_data: Any, input_type: str):
        """Validate input before processing"""
        if input_type == 'text':
            if not isinstance(input_data, str):
                raise TypeError("Text input must be a string")
            if len(input_data) < 5:
                raise ValueError("Text is too short (minimum 5 characters)")

        elif input_type in ['voice', 'pdf', 'docx', 'xlsx', 'csv', 'image']:
            if not os.path.exists(input_data):
                raise FileNotFoundError(f"File not found: {input_data}")

            # Check file size (max 100MB)
            file_size = os.path.getsize(input_data)
            max_size = 100 * 1024 * 1024  # 100MB
            if file_size > max_size:
                raise ValueError(f"File too large: {file_size / 1024 / 1024:.1f}MB (max 100MB)")

        elif input_type == 'url':
            if not isinstance(input_data, str):
                raise TypeError("URL must be a string")
            if not input_data.startswith(('http://', 'https://')):
                raise ValueError("Invalid URL format")


# ==========================================
# TESTING
# ==========================================

if __name__ == "__main__":
    # Initialize processor
    processor = InputProcessor(whisper_model_size="base")
    
    # Test 1: Text input
    print("\n" + "="*60)
    print("TEST 1: Text Input")
    print("="*60)

    text_input = "India's GDP grew 8% in 2024 according to official government data released yesterday."
    result = processor.process(text_input, 'text')

    print(f"Text: {result['text']}")
    print(f"Language: {result['language']}")
    print(f"Metadata: {json.dumps(result['metadata'], indent=2)}")

    # Test 2: URL
    print("\n" + "="*60)
    print("TEST 2: URL Input")
    print("="*60)

    url = "https://www.bbc.com/news/world"
    try:
        result = processor.process(url, 'url')
        print(f"Title: {result['metadata']['title']}")
        print(f"Text preview: {result['text'][:200]}...")
        print(f"Language: {result['language']}")
    except Exception as e:
        print(f"Error: {str(e)}")