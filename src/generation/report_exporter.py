"""
src/generation/report_exporter.py
───────────────────────────────────
Exports the LLM-generated markdown report to:
  - HTML  : dark-themed, verdict badge, confidence bar, evidence stance chart, credibility chart
  - PDF   : via weasyprint from the styled HTML
  - DOCX  : clean Word document with formatting
  - MD    : raw markdown

The HTML report is the primary output — visually rich, printable, and standalone.
"""
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any


class ReportExporter:

    def __init__(self, output_dir: str = "data/reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _filename(self, claim: str, ext: str) -> Path:
        ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = re.sub(r"[^\w\s-]", "", claim[:40]).strip().replace(" ", "_")
        return self.output_dir / f"report_{safe}_{ts}.{ext}"

    # ── Markdown ──────────────────────────────────────────────────────────────

    def to_markdown(self, report_md: str, claim: str = "claim") -> str:
        path = self._filename(claim, "md")
        path.write_text(report_md, encoding="utf-8")
        return str(path)

    # ── HTML (beautiful) ──────────────────────────────────────────────────────

    def to_html(self, report_md: str, claim: str = "claim",
                rag_result: Optional[Dict[str, Any]] = None) -> str:
        """Convert report to fully styled HTML with charts and verdict badge."""

        # Convert markdown to HTML
        try:
            import markdown as md_lib
            body_html = md_lib.markdown(report_md,
                                         extensions=["tables", "fenced_code", "nl2br"])
        except ImportError:
            body_html = "<pre style='white-space:pre-wrap'>" + report_md + "</pre>"

        # Parse verdict and confidence from report text
        verdict    = self._parse_verdict(report_md, rag_result)
        confidence = self._parse_confidence(report_md, rag_result)

        # Parse evidence counts from report text
        supports = len(re.findall(r"SUPPORTS", report_md, re.I))
        refutes  = len(re.findall(r"REFUTES",  report_md, re.I))
        neutral  = max(0, len(re.findall(r"NEUTRAL",  report_md, re.I)) - 1)
        total_ev = max(supports + refutes + neutral, 1)
        sup_pct  = round(supports / total_ev * 100)
        ref_pct  = round(refutes  / total_ev * 100)
        neu_pct  = 100 - sup_pct - ref_pct

        # If rag_result provided use its actual counts
        if rag_result and rag_result.get("aggregation"):
            agg      = rag_result["aggregation"]
            supports = agg.get("num_supports", supports)
            refutes  = agg.get("num_refutes",  refutes)
            neutral  = agg.get("num_neutral",  neutral)
            total_ev = max(supports + refutes + neutral, 1)
            sup_pct  = round(supports / total_ev * 100)
            ref_pct  = round(refutes  / total_ev * 100)
            neu_pct  = 100 - sup_pct - ref_pct

        # Verdict colour
        vc = {"TRUE": "#34d399", "MOSTLY TRUE": "#6ee7b7",
              "UNVERIFIABLE": "#fbbf24", "MOSTLY FALSE": "#fb923c",
              "FALSE": "#f87171", "CONFLICTING": "#a78bfa"}.get(verdict, "#94a3b8")
        vb = {"TRUE": "rgba(52,211,153,0.12)", "MOSTLY TRUE": "rgba(110,231,183,0.10)",
              "UNVERIFIABLE": "rgba(251,191,36,0.12)", "MOSTLY FALSE": "rgba(251,146,60,0.12)",
              "FALSE": "rgba(248,113,113,0.12)", "CONFLICTING": "rgba(167,139,250,0.12)"
              }.get(verdict, "rgba(148,163,184,0.10)")

        ts = datetime.now().strftime("%B %d, %Y at %H:%M")

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Fact Verification Report</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Mono:wght@400;500&family=DM+Sans:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
:root{{
  --bg:#0d0f14;--s:#13161e;--s2:#1a1e2a;--b:#252836;
  --text:#e2e8f0;--muted:#64748b;--sub:#94a3b8;
  --green:#34d399;--amber:#fbbf24;--red:#f87171;--blue:#6c8cff;
  --serif:'DM Serif Display',Georgia,serif;
  --mono:'DM Mono',monospace;--sans:'DM Sans',sans-serif;
}}
body{{background:var(--bg);color:var(--text);font-family:var(--sans);
      font-size:15px;line-height:1.75;-webkit-font-smoothing:antialiased}}
.wrap{{max-width:880px;margin:0 auto;padding:48px 32px 80px}}
.topbar{{display:flex;align-items:center;justify-content:space-between;
         margin-bottom:48px;padding-bottom:20px;border-bottom:1px solid var(--b)}}
.logo{{font-family:var(--serif);font-size:1.3rem;color:var(--blue);font-style:italic}}
.ts{{font-family:var(--mono);font-size:.72rem;color:var(--muted);letter-spacing:.06em}}
.claim-block{{background:var(--s);border:1px solid var(--b);border-radius:12px;
              padding:28px 32px;margin-bottom:28px}}
.label{{font-family:var(--mono);font-size:.68rem;letter-spacing:.14em;
        text-transform:uppercase;color:var(--muted);margin-bottom:10px}}
.claim-text{{font-family:var(--serif);font-size:1.45rem;line-height:1.4;
             color:var(--text);font-style:italic}}
.verdict-card{{border-radius:12px;padding:28px 32px;margin-bottom:28px;
               border:1px solid {vc}44;background:{vb}}}
.verdict-row{{display:flex;align-items:center;gap:20px;flex-wrap:wrap}}
.verdict-badge{{font-family:var(--mono);font-size:1.15rem;font-weight:500;
                color:{vc};letter-spacing:.04em}}
.conf-wrap{{flex:1;min-width:180px}}
.conf-label{{font-family:var(--mono);font-size:.7rem;color:var(--muted);margin-bottom:6px}}
.conf-bar{{height:7px;background:var(--b);border-radius:4px;overflow:hidden}}
.conf-fill{{height:100%;border-radius:4px;background:{vc};width:{confidence}%}}
.conf-pct{{font-family:var(--mono);font-size:.9rem;color:{vc};margin-top:5px}}
.charts{{display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:28px}}
.chart-card{{background:var(--s);border:1px solid var(--b);border-radius:12px;padding:24px}}
.chart-title{{font-family:var(--mono);font-size:.68rem;letter-spacing:.12em;
              text-transform:uppercase;color:var(--muted);margin-bottom:16px}}
.bar-row{{display:flex;align-items:center;gap:10px;margin-bottom:10px}}
.bar-label{{font-family:var(--mono);font-size:.72rem;color:var(--sub);width:72px}}
.bar-track{{flex:1;height:8px;background:var(--b);border-radius:4px;overflow:hidden}}
.bar-fill{{height:100%;border-radius:4px}}
.bar-pct{{font-family:var(--mono);font-size:.72rem;color:var(--muted);width:34px;text-align:right}}
.sum-row{{display:flex;align-items:center;justify-content:space-between;
          padding:8px 0;border-bottom:1px solid var(--b)}}
.sum-row:last-child{{border-bottom:none}}
.sum-name{{font-family:var(--mono);font-size:.78rem;color:var(--sub)}}
.sum-val{{font-family:var(--mono);font-size:.78rem}}
.report-body{{background:var(--s);border:1px solid var(--b);border-radius:12px;
              padding:36px 40px;margin-bottom:28px}}
.report-body h1,.report-body h2{{font-family:var(--serif);color:var(--text);
  margin-top:32px;margin-bottom:12px;line-height:1.3}}
.report-body h1{{font-size:1.6rem;font-style:italic;
  border-bottom:1px solid var(--b);padding-bottom:12px}}
.report-body h2{{font-size:1.2rem}}
.report-body h3{{font-size:1rem;color:var(--sub);margin-top:20px;margin-bottom:8px}}
.report-body p{{margin-bottom:14px;color:var(--sub)}}
.report-body strong{{color:var(--text);font-weight:600}}
.report-body ul,.report-body ol{{padding-left:20px;margin-bottom:14px;color:var(--sub)}}
.report-body li{{margin-bottom:6px}}
.report-body table{{width:100%;border-collapse:collapse;margin-bottom:20px;font-size:.88rem}}
.report-body th{{background:var(--s2);color:var(--text);padding:10px 14px;text-align:left;
  font-family:var(--mono);font-size:.72rem;letter-spacing:.08em;text-transform:uppercase}}
.report-body td{{padding:10px 14px;border-bottom:1px solid var(--b);color:var(--sub)}}
.report-body code{{background:var(--s2);padding:2px 6px;border-radius:4px;
  font-family:var(--mono);font-size:.85em;color:var(--blue)}}
.report-body blockquote{{border-left:3px solid var(--blue);padding:10px 18px;
  margin:16px 0;background:var(--s2);border-radius:0 6px 6px 0;
  color:var(--sub);font-style:italic}}
.footer{{text-align:center;padding-top:32px;border-top:1px solid var(--b)}}
.footer p{{font-family:var(--mono);font-size:.7rem;color:var(--muted);letter-spacing:.06em}}
@media print{{
  body{{background:#fff!important;color:#000!important}}
  .charts{{display:none!important}}
  .report-body{{background:#fff!important;border:none!important;padding:0!important}}
  .claim-block,.verdict-card{{background:#f8f9fa!important;border:1px solid #ddd!important}}
}}
</style>
</head>
<body>
<div class="wrap">

<div class="topbar">
  <span class="logo">FactCheck AI</span>
  <span class="ts">Generated {ts}</span>
</div>

<div class="claim-block">
  <div class="label">Claim under verification</div>
  <div class="claim-text">&#8220;{claim}&#8221;</div>
</div>

<div class="verdict-card">
  <div class="verdict-row">
    <div>
      <div class="label">Verdict</div>
      <div class="verdict-badge">{verdict}</div>
    </div>
    <div class="conf-wrap">
      <div class="conf-label">Confidence score</div>
      <div class="conf-bar"><div class="conf-fill"></div></div>
      <div class="conf-pct">{confidence}%</div>
    </div>
  </div>
</div>

<div class="charts">
  <div class="chart-card">
    <div class="chart-title">Evidence stance breakdown</div>
    <div class="bar-row">
      <span class="bar-label">Supports</span>
      <div class="bar-track"><div class="bar-fill" style="width:{sup_pct}%;background:#34d399"></div></div>
      <span class="bar-pct">{sup_pct}%</span>
    </div>
    <div class="bar-row">
      <span class="bar-label">Refutes</span>
      <div class="bar-track"><div class="bar-fill" style="width:{ref_pct}%;background:#f87171"></div></div>
      <span class="bar-pct">{ref_pct}%</span>
    </div>
    <div class="bar-row">
      <span class="bar-label">Neutral</span>
      <div class="bar-track"><div class="bar-fill" style="width:{neu_pct}%;background:#6c8cff"></div></div>
      <span class="bar-pct">{neu_pct}%</span>
    </div>
  </div>
  <div class="chart-card">
    <div class="chart-title">Verification summary</div>
    <div class="sum-row"><span class="sum-name">Total evidence</span>
      <span class="sum-val" style="color:var(--blue)">{supports+refutes+neutral} pieces</span></div>
    <div class="sum-row"><span class="sum-name">Supporting</span>
      <span class="sum-val" style="color:var(--green)">{supports}</span></div>
    <div class="sum-row"><span class="sum-name">Refuting</span>
      <span class="sum-val" style="color:var(--red)">{refutes}</span></div>
    <div class="sum-row"><span class="sum-name">Neutral</span>
      <span class="sum-val" style="color:var(--muted)">{neutral}</span></div>
    <div class="sum-row"><span class="sum-name">Confidence</span>
      <span class="sum-val" style="color:{vc}">{confidence}%</span></div>
  </div>
</div>

<div class="report-body">{body_html}</div>

<div class="footer">
  <p>FACTCHECK AI &nbsp;&middot;&nbsp; EVIDENCE-GROUNDED VERIFICATION &nbsp;&middot;&nbsp; {ts.upper()}</p>
</div>

</div>
</body>
</html>"""

        path = self._filename(claim, "html")
        path.write_text(html, encoding="utf-8")
        return str(path)

    # ── PDF ───────────────────────────────────────────────────────────────────

    def to_pdf(self, report_md: str, claim: str = "claim",
               rag_result: Optional[Dict[str, Any]] = None) -> str:
        try:
            from weasyprint import HTML as WP
        except ImportError:
            raise ImportError("pip install weasyprint")
        html_path = self.to_html(report_md, claim, rag_result)
        pdf_path  = self._filename(claim, "pdf")
        WP(filename=html_path).write_pdf(str(pdf_path))
        return str(pdf_path)

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def to_docx(self, report_md: str, claim: str = "claim",
                rag_result: Optional[Dict[str, Any]] = None) -> str:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("pip install python-docx")

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Pt(64)
            section.bottom_margin = Pt(64)
            section.left_margin   = Pt(80)
            section.right_margin  = Pt(80)

        # Title
        t = doc.add_heading("Fact Verification Report", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Claim: {claim}")
        verdict    = self._parse_verdict(report_md, rag_result)
        confidence = self._parse_confidence(report_md, rag_result)
        doc.add_paragraph(f"Verdict: {verdict}  |  Confidence: {confidence}%")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        for line in report_md.splitlines():
            s = line.strip()
            if not s:
                doc.add_paragraph()
            elif s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s.startswith(("- ", "* ")):
                doc.add_paragraph(s[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", s):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", s), style="List Number")
            else:
                para  = doc.add_paragraph()
                parts = re.split(r"\*\*(.+?)\*\*", s)
                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    if i % 2 == 1:
                        run.bold = True

        doc.add_paragraph()
        footer = doc.add_paragraph("Powered by FactCheck AI")
        footer.runs[0].italic = True
        footer.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        path = self._filename(claim, "docx")
        doc.save(str(path))
        return str(path)

    # ── Export all ────────────────────────────────────────────────────────────

    def export_all(self, report_md: str, claim: str = "claim",
                   rag_result: Optional[Dict[str, Any]] = None) -> dict:
        results = {}
        results["markdown"] = self.to_markdown(report_md, claim)
        results["html"]     = self.to_html(report_md, claim, rag_result)
        try:
            results["pdf"]  = self.to_pdf(report_md, claim, rag_result)
        except Exception as e:
            results["pdf"]  = f"SKIPPED: {e}"
        try:
            results["docx"] = self.to_docx(report_md, claim, rag_result)
        except Exception as e:
            results["docx"] = f"SKIPPED: {e}"
        return results

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _parse_verdict(self, report_md: str,
                       rag_result: Optional[Dict] = None) -> str:
        if rag_result and rag_result.get("verdict"):
            return rag_result["verdict"]
        for line in report_md.splitlines():
            lu = line.upper()
            if "MOSTLY TRUE"  in lu: return "MOSTLY TRUE"
            if "MOSTLY FALSE" in lu: return "MOSTLY FALSE"
            if "UNVERIFIABLE" in lu: return "UNVERIFIABLE"
            if "CONFLICTING"  in lu: return "CONFLICTING"
            if "**TRUE**"     in lu or "VERDICT: TRUE"  in lu: return "TRUE"
            if "**FALSE**"    in lu or "VERDICT: FALSE" in lu: return "FALSE"
        return "UNVERIFIABLE"

    def _parse_confidence(self, report_md: str,
                          rag_result: Optional[Dict] = None) -> float:
        if rag_result and rag_result.get("confidence") is not None:
            return round(float(rag_result["confidence"]), 1)
        m = re.search(r"[Cc]onfidence[:\s]+([0-9.]+)%", report_md)
        if m:
            return round(float(m.group(1)), 1)
        return 0.0